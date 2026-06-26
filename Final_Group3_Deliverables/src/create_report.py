from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import pandas as pd

OUT = Path('/mnt/data/Amazon_Enterprise_Risk_Assessment_Report.docx')
CHART_DIR = Path('/mnt/data/risk_outputs')

def clean(s):
    return str(s).replace('-', ' ').replace('–', ' ').replace('—', ' ')

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=8.5):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(clean(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP



def set_cell_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in('w:tcW')
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)

def add_para(doc, text='', style=None, align=None, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(clean(text))
        r.bold = bold
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f'Heading {level}' if level <= 3 else 'Normal'
    r = p.add_run(clean(text))
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14 if level == 1 else 12)
    if level == 1:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows, widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    hdr = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    set_cant_split(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=font_size)
        set_cell_shading(hdr[i], 'D9EAF7')
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row in rows:
        row_obj = table.add_row()
        set_cant_split(row_obj)
        cells = row_obj.cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if widths:
                set_cell_width(cells[i], widths[i])
    add_para(doc, '')
    return table

# Load computed model results
qdf = pd.read_csv(CHART_DIR / 'quantitative_risk_summary.csv')
ml_metrics = pd.read_csv(CHART_DIR / 'ml_model_metrics.csv')
ml_coef = pd.read_csv(CHART_DIR / 'ml_logistic_coefficients.csv')
tail = pd.read_csv(CHART_DIR / 'tail_contribution.csv')

# Create document
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

# Base font and paragraph spacing
styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal'].font.size = Pt(11)
for style_name in ['Normal']:
    style = styles[style_name]
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(4)

for h in ['Heading 1','Heading 2','Heading 3']:
    styles[h].font.name = 'Times New Roman'
    styles[h].font.color.rgb = RGBColor(31, 78, 121)
    styles[h].paragraph_format.space_after = Pt(4)

# Title page
p = add_para(doc, 'Enterprise Risk Management Analytics Assessment for Amazon AI Warehouse and Customer Service Strategy', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
p.runs[0].font.size = Pt(18)
add_para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'ALY6130: Risk Management for Analytics', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, 'Signature Assessment Report', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'Prepared for decision makers', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'Group 3', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'June 2026', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '')
add_para(doc, 'Project note: This report builds from the earlier Amazon SWOT analysis and updates the risk register so that it uses only the ten risks requested by the lecturer. The analysis moves beyond description by adding probability distributions, expected loss estimates, machine learning based prediction using a synthetic proxy dataset, Monte Carlo simulation, and KRI trigger points for action.', align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

add_heading(doc, 'Summary for Decision Makers', 1)
add_para(doc, 'Amazon has the financial strength, cloud capability, and operational data needed to benefit from AI in customer service and warehouse operations. At the same time, the same strengths create serious enterprise risks because AI decisions now touch customer trust, fulfillment speed, employee adoption, privacy obligations, and the cost of technology investment. The earlier SWOT work identified technology dependence, cybersecurity exposure, workforce adaptation, competitive pressure, and regulatory complexity as key areas of concern. This report converts those concerns into a structured enterprise risk management assessment using the ten risks selected by the lecturer.')
add_para(doc, 'The quantitative model estimates a base case expected annual exposure of about 696 million dollars across the ten risks. The Monte Carlo simulation, which allows risks to occur together through market, technology, regulatory, and people factors, produces a mean annual loss of about 719 million dollars and a 95th percentile annual loss of about 1.59 billion dollars. This does not mean Amazon will lose that amount. It means that, under the assumptions used in this model, decision makers should plan controls and contingency funding for a much wider loss range than the simple average suggests.')
add_para(doc, 'The largest expected annual loss in the model comes from cost overrun and delayed ROI, regulatory and privacy non compliance, and cybersecurity and data breach risk. The largest tail contributors are cybersecurity, regulatory non compliance, and cost overrun. For management action, this means the risk response should not only focus on technical implementation. It should also give equal attention to governance, compliance, privacy design, cloud resilience, training, and budget discipline.')

add_heading(doc, 'Background', 1)
add_para(doc, 'Amazon is a global technology, retail, cloud computing, logistics, advertising, and digital services company. In 2025, Amazon reported net sales of 716.9 billion dollars, while AWS sales reached 128.7 billion dollars. Operating income increased to 80.0 billion dollars, and the company stated that its free cash flow decreased partly because property and equipment purchases increased due to artificial intelligence investments (Amazon, 2026a). These figures matter for risk management because they show that Amazon can fund major AI programs, but they also show that AI investment has become large enough to affect enterprise financial exposure.')
add_para(doc, 'The business scenario for this assessment is Amazon acquiring and integrating an AI powered customer service platform while also expanding AI capabilities into warehouse strategy, forecasting, robotics, cloud enabled operations, and governance. The earlier SWOT analysis for this project argued that Amazon has advanced technology infrastructure, strong financial capacity, and a large global customer base. It also identified weaknesses in technology dependence, cybersecurity and data privacy, and employee resistance (Group 3, 2026). Those points remain important, but the final assessment must go further by quantifying the risks and connecting them to measurable KRIs.')
add_para(doc, 'Amazon already operates at a scale where AI risk is not theoretical. Its 2025 annual report states that competition is intense across retail, cloud, fulfillment, logistics, customer service, advertising, and artificial intelligence services. The same filing also notes that AI and machine learning are increasing competition, that technology and infrastructure spending is expected to increase, and that system interruptions can affect order fulfillment and services (Amazon, 2026c). This creates a clear risk environment: AI can improve Amazon operations, but AI can also increase the speed, scale, and visibility of failure.')
add_para(doc, 'Amazon has also continued to expand robotics and AI in fulfillment. In 2025, Amazon reported that it deployed its one millionth robot and introduced DeepFleet, a generative AI foundation model designed to improve robot fleet travel efficiency by 10 percent. Amazon also reported that more than 700,000 employees had been upskilled for the future (Amazon, 2025). In 2026, Amazon announced further robotics expansion in Europe, including more than 10 billion euros in fulfillment center modernization and 1 billion dollars in employee upskilling (Amazon, 2026b). These details directly support the focus on AI warehouse strategy failure, workforce adoption failure, cost overrun, and governance risk.')
add_para(doc, 'The risk management structure used in this report aligns with ISO 31000, which presents risk management as a principles based process for identifying, analyzing, evaluating, and treating risk (International Organization for Standardization, 2018). The AI governance portion also aligns with the NIST AI Risk Management Framework, which focuses on Govern, Map, Measure, and Manage functions for trustworthy AI (National Institute of Standards and Technology, 2023). These standards help turn the lecturer approved risk list into an evidence based ERM assessment that decision makers can use.')

add_heading(doc, 'Identification of Risks', 1)
add_para(doc, 'The ten risks below are retained in the order given by the lecturer. The order is treated as the official identification order, while the later qualitative and quantitative scores provide additional insight about priority, exposure, and trigger points. The risks are linked to the earlier SWOT findings because they flow from the same business context: AI driven competition, warehouse automation, customer data processing, cloud dependence, employee adaptation, regulatory pressure, and technology investment.')

ident_rows = [
['1', 'Competitor AI Advantage', 'Strategic', 'Rivals use superior AI tools in customer service, logistics, cloud services, or retail personalization before Amazon can respond.', 'Loss of differentiation, lower customer retention, and pressure to spend faster.'],
['2', 'AI Warehouse Strategy Failure', 'Strategic and operational', 'AI robotics and warehouse programs do not achieve expected throughput, travel efficiency, safety, or process redesign benefits.', 'Reduced productivity, delayed fulfillment improvements, and lower strategic return.'],
['3', 'Cybersecurity & Data Breach', 'Information security', 'The acquired and integrated AI platform processes sensitive customer, transaction, and support data that could be exposed or misused.', 'Legal cost, customer trust erosion, regulatory scrutiny, and disruption.'],
['4', 'AI Forecasting Failure', 'Technology and operational', 'Forecasting models fail because demand patterns change, training data quality is weak, or algorithms drift.', 'Stockouts, overstocking, poor labor planning, and increased fulfillment cost.'],
['5', 'System Integration Failure', 'Technology and operational', 'Legacy systems, customer service systems, warehouse platforms, APIs, and data pipelines do not connect reliably.', 'Process interruption, inconsistent customer records, and delayed implementation.'],
['6', 'Cloud Service Outage', 'Technology and continuity', 'Cloud infrastructure or dependent services become unavailable or slow during critical operations.', 'Customer service delays, fulfillment disruption, and lost revenue.'],
['7', 'Workforce Adoption Failure', 'People and change', 'Employees do not trust, understand, or consistently use AI tools because training and change management are weak.', 'Low adoption, resistance, morale concerns, and reduced productivity.'],
['8', 'Cost Overrun & Delayed ROI', 'Financial', 'AI integration, cloud capacity, vendor costs, training, data work, and governance costs exceed budget or benefits arrive later than planned.', 'Reduced profitability, investment pressure, and weaker confidence in the strategy.'],
['9', 'Regulatory & Privacy Non Compliance', 'Compliance', 'Privacy, data security, consumer protection, and AI regulation requirements are not met across jurisdictions.', 'Fines, remediation cost, product restrictions, and reputational damage.'],
['10', 'AI Governance & Ethics Failure', 'Governance and compliance', 'AI systems produce biased, unfair, opaque, or poorly controlled outcomes without adequate oversight.', 'Loss of trust, internal control weakness, regulatory action, and ethical concern.'],
]
add_table(doc, ['Rank', 'Risk', 'Category', 'Core cause', 'Potential enterprise impact'], ident_rows, widths=[0.4,1.3,1.0,2.4,2.0], font_size=7.6)

add_heading(doc, 'Qualitative Risk Assessment', 1)
add_para(doc, 'The qualitative assessment uses a five point scale for likelihood and impact. A score of 1 means low, 3 means moderate, and 5 means very high. The score is calculated as likelihood multiplied by impact. This is useful for quick decision making, but it is not enough for a final ERM assessment because it does not show loss variation, tail exposure, or the effect of several risks happening together. For that reason, the qualitative results are used as inputs to the quantitative model rather than as the final answer.')
qual_rows = [
['1', 'Competitor AI Advantage', '4', '5', '20', 'Critical'],
['2', 'AI Warehouse Strategy Failure', '4', '5', '20', 'Critical'],
['3', 'Cybersecurity & Data Breach', '3', '5', '15', 'High'],
['4', 'AI Forecasting Failure', '4', '4', '16', 'High'],
['5', 'System Integration Failure', '3', '4', '12', 'High'],
['6', 'Cloud Service Outage', '3', '4', '12', 'High'],
['7', 'Workforce Adoption Failure', '4', '3', '12', 'High'],
['8', 'Cost Overrun & Delayed ROI', '4', '4', '16', 'High'],
['9', 'Regulatory & Privacy Non Compliance', '3', '5', '15', 'High'],
['10', 'AI Governance & Ethics Failure', '3', '4', '12', 'High'],
]
add_table(doc, ['Rank', 'Risk', 'Likelihood', 'Impact', 'Score', 'Rating'], qual_rows, widths=[0.45,2.45,0.9,0.75,0.7,0.9], font_size=8.0)
add_para(doc, 'Competitor AI advantage and AI warehouse strategy failure are rated critical because they directly affect Amazon strategy and market position. The annual report confirms that AI and machine learning are increasing competitive pressure, and the scale of robotics expansion means a weak warehouse strategy could affect productivity, fulfillment speed, and expected automation value. Cybersecurity, regulatory non compliance, and cost overrun are high because their impacts can become severe even when likelihood is moderate. Workforce adoption and governance are also high because weak adoption, poor documentation, or low trust can weaken even a technically strong implementation.')
add_heading(doc, 'Quantitative Risk Assessment', 1)
add_para(doc, 'The quantitative assessment applies three methods required for the Signature Assessment. The first method is standardized quantitative risk modeling using probability distributions and expected loss. The second method is machine learning based risk prediction using a synthetic proxy dataset that reflects Amazon style operating conditions. The third method is Monte Carlo simulation to model uncertainty, correlation, and enterprise level exposure. Because Amazon internal incident data is not publicly available, the assumptions are transparent and intentionally conservative. The figures should be read as decision support estimates, not as actual Amazon loss history.')

add_heading(doc, 'Standardized Quantitative Risk Modeling', 2)
add_para(doc, 'Each risk is assigned a beta probability distribution for annual likelihood and a triangular distribution for financial impact. The beta distribution is useful because it keeps probabilities between 0 and 1 and allows expert judgment to express uncertainty. The triangular distribution is useful when there is no internal dataset but there is a reasonable low, most likely, and high impact estimate. Expected loss is calculated as mean probability multiplied by expected impact. Variance and 95th percentile loss are estimated from repeated sampling.')
add_para(doc, 'The financial impact values are in millions of dollars. The estimates reflect business scale, likely remediation cost, opportunity cost, and the severity implied by the earlier SWOT findings. For example, cybersecurity and regulatory risks receive high impact ranges because breach response, legal exposure, and trust loss can extend beyond direct technical repair. Cost overrun receives a high expected loss because Amazon has already signaled heavy technology and infrastructure investment connected to AI and cloud growth (Amazon, 2026a; Amazon, 2026c).')
q_rows=[]
for _, r in qdf.iterrows():
    q_rows.append([r['Risk'], r['Risk Name'], f"{r['Mean Probability']:.2f}", f"{r['Expected Impact $m']:.1f}", f"{r['Expected Loss $m']:.1f}", f"{r['Loss Variance']:.1f}", f"{r['P95 Loss $m']:.1f}"])
add_table(doc, ['Risk', 'Risk name', 'Mean probability', 'Expected impact $m', 'Expected loss $m', 'Variance', '95 percentile loss $m'], q_rows, widths=[0.55,1.55,0.8,0.95,0.95,0.75,0.95], font_size=7.3)
add_para(doc, 'The total expected annual exposure across the ten risks is about 696 million dollars. The highest expected losses are cost overrun and delayed ROI at 113.2 million dollars, regulatory and privacy non compliance at 104.0 million dollars, and cybersecurity and data breach at 101.9 million dollars. This does not mean these are the only important risks. Rather, they combine meaningful probability with large potential impact. Competitor AI advantage and AI warehouse strategy failure remain strategically important even when their expected loss is lower than some compliance or financial exposures.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(str(CHART_DIR / 'expected_loss_bar.png'), width=Inches(6.2))
add_para(doc, 'Figure 1. Expected annual loss by risk. Source: synthetic quantitative model developed for this assessment.', align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading(doc, 'Scenario Based Numerical Modeling', 2)
add_para(doc, 'The scenario model uses three views of the business and non business environment. The base case assumes current competitive pressure, current AI investment levels, and normal regulatory conditions. The stress case assumes stronger competitor activity, higher AI infrastructure cost, more integration defects, and tighter privacy expectations. The severe case assumes simultaneous technology disruption, market pressure, regulatory attention, and workforce adoption problems. These scenarios connect directly to the assignment requirement to consider the business environment and the non business environment together.')
scenario_rows = [
['Base case', 'Current AI investment and competitive pressure continue. Controls are implemented but still maturing.', '1.00', '1.00', '696.0', '1,585.9'],
['Stress case', 'AI cost pressure rises, warehouse benefits are delayed, and regulators increase scrutiny.', '1.25', '1.15', '1,000.5', '2,279.7'],
['Severe case', 'Competitor gains, cloud instability, privacy events, and governance gaps occur in the same planning period.', '1.55', '1.35', '1,456.4', '3,318.5'],
]
add_table(doc, ['Scenario', 'Assumption', 'Probability multiplier', 'Impact multiplier', 'Expected exposure $m', 'Modelled 95 percentile $m'], scenario_rows, widths=[0.85,3.1,0.75,0.75,0.8,0.9], font_size=7.5)
add_para(doc, 'The scenario results show why a static risk register is not enough. If Amazon manages each risk separately, it could underestimate exposure during a period when market competition, regulation, cloud capacity, and workforce pressure move together. This is especially relevant because Amazon stated that macroeconomic factors, tariffs, geopolitical developments, and AI adoption can affect forecasting, expenses, and technology benefits in ways that are difficult to predict and quantify (Amazon, 2026c).')

add_heading(doc, 'ML Based Risk Prediction', 2)
add_para(doc, 'A supervised learning model was developed using a synthetic proxy dataset of 1,200 monthly observations. Each observation represents a business unit, warehouse program, or service operation month. The target variable is a high loss risk month, defined as a month where combined KRI signals imply elevated likelihood of a significant loss event. The dataset is synthetic because Amazon internal incident data is not publicly available. It was generated to reflect realistic operational signals: competitor AI release index, warehouse productivity gap, security incident count, forecast error, integration defects, cloud downtime, employee adoption, budget variance, compliance gaps, bias reports, and training completion.')
add_para(doc, 'Feature selection was aligned with the identified KRIs. The first model was logistic regression because it is easier to explain to decision makers. A random forest model was also tested as a nonlinear benchmark. The data was divided into 70 percent training and 30 percent validation sets, with stratification so that high loss months appeared in both sets. The logistic regression model produced the stronger validation result and is recommended for reporting because its coefficients are transparent.')
ml_rows=[]
for _, r in ml_metrics.iterrows():
    ml_rows.append([r['Model'], f"{r['ROC AUC']:.3f}", f"{r['Accuracy']:.3f}", f"{r['Precision']:.3f}", f"{r['Recall']:.3f}", f"{r['F1']:.3f}"])
add_table(doc, ['Model', 'ROC AUC', 'Accuracy', 'Precision', 'Recall', 'F1'], ml_rows, widths=[1.6,0.8,0.8,0.8,0.8,0.8], font_size=8.0)
coef_rows=[]
for _, r in ml_coef.head(8).iterrows():
    direction = 'Increases predicted risk' if r['Standardized coefficient'] > 0 else 'Reduces predicted risk'
    coef_rows.append([r['Feature'].replace('_',' '), f"{r['Standardized coefficient']:.3f}", direction])
add_table(doc, ['Feature', 'Standardized coefficient', 'Interpretation'], coef_rows, widths=[2.2,1.1,2.8], font_size=8.0)
add_para(doc, 'The strongest predictors in the logistic model are budget variance, forecast error, open compliance gaps, employee adoption, warehouse productivity gap, bias reports, security incidents, and cloud downtime. This output is useful because it agrees with the qualitative analysis. It shows that financial control, model performance, compliance status, people adoption, and technology continuity should be monitored together rather than in separate dashboards. Employee adoption has a negative coefficient, meaning that higher adoption reduces predicted risk. Training completion also reduces predicted risk, although it is weaker than actual adoption.')

add_heading(doc, 'Monte Carlo Simulation', 2)
add_para(doc, 'The Monte Carlo simulation ran 100,000 annual trials. In each trial, the model sampled whether each risk occurred and then sampled the loss amount if the event occurred. The simulation included common market, technology, regulatory, and people factors so that risks could move together. This matters because cybersecurity, cloud outage, integration failure, and governance failure may all worsen in the same period if technology complexity rises. Likewise, competitor pressure and cost overrun can increase together if Amazon accelerates spending to protect market position.')
mc_rows = [
['Mean annual loss', '718.7 million dollars'],
['Median annual loss', '654.1 million dollars'],
['75 percentile annual loss', '1,002.9 million dollars'],
['90 percentile annual loss', '1,356.9 million dollars'],
['95 percentile annual loss', '1,585.9 million dollars'],
['99 percentile annual loss', '2,046.6 million dollars'],
['Probability annual loss exceeds 500 million dollars', '63.3 percent'],
['Probability annual loss exceeds 1 billion dollars', '25.2 percent'],
]
add_table(doc, ['Monte Carlo output', 'Result'], mc_rows, widths=[3.2,2.4], font_size=8.2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(str(CHART_DIR / 'monte_carlo_loss_distribution.png'), width=Inches(6.2))
add_para(doc, 'Figure 2. Monte Carlo enterprise loss distribution. Source: synthetic simulation developed for this assessment.', align=WD_ALIGN_PARAGRAPH.CENTER)

tail_rows=[]
for _, r in tail.sort_values('Tail contribution', ascending=False).head(5).iterrows():
    tail_rows.append([r['Risk'], f"{r['Tail contribution']*100:.1f} percent"])
add_table(doc, ['Largest contributors in 95 percentile loss years', 'Share of tail exposure'], tail_rows, widths=[3.5,1.5], font_size=8.2)
add_para(doc, 'The tail contribution result is important for decision makers because it shows where the worst years are coming from. Cybersecurity and data breach contributes about 22.4 percent of 95th percentile years, regulatory and privacy non compliance contributes about 20.1 percent, and cost overrun and delayed ROI contributes about 14.2 percent. Therefore, the most practical capital protection strategy is to strengthen security controls, privacy compliance, and investment governance before the AI program scales further.')

add_heading(doc, 'Integrated Situation Assessment', 2)
add_para(doc, 'The business environment influences the model inputs in several ways. Market competition increases the probability of competitor AI advantage and cost pressure. Amazon’s financial strength reduces the likelihood of running out of funds, but it does not remove budget variance or delayed ROI because large AI infrastructure investments can grow faster than expected benefits. Supply chain and fulfillment complexity increase the impact of AI warehouse strategy failure, forecasting failure, and system integration failure. If forecasting quality declines, the impact spreads to inventory placement, staffing, transportation cost, and customer promise reliability.')
add_para(doc, 'The non business environment also changes the risk distributions. Regulation increases the impact range for privacy non compliance and AI governance failure. The GDPR allows severe administrative fines up to 20 million euros or 4 percent of total worldwide annual turnover, whichever is higher (European Parliament and Council of the European Union, 2016). The EU AI Act has also created a more structured legal environment for AI governance and trustworthy AI (European Commission, 2024). Cybersecurity conditions increase both probability and impact because AI systems introduce new data flows, APIs, model access points, and third party dependencies. IBM reported a global average data breach cost of 4.44 million dollars in 2025 and highlighted governance gaps around AI systems (IBM, 2025).')
add_para(doc, 'The technological environment is the strongest shared factor across the model. Cloud service outage, integration failure, forecasting failure, and AI governance failure all rely on data quality, resilient architecture, monitoring, and change control. The AWS Well Architected Reliability Pillar emphasizes workload reliability, recovery from failure, distributed design, change management, and monitoring (Amazon Web Services, 2024). For Amazon, this means the risk model inputs should be updated whenever major cloud architecture, AI model, vendor, or warehouse process changes occur.')

add_heading(doc, 'Risk Response Strategy', 1)
add_para(doc, 'The response strategy combines mitigation, transfer, acceptance, and contingency planning. The main recommendation is not to slow down AI adoption. The better approach is to scale AI with stronger controls so that Amazon can protect value while still using AI as a competitive advantage. The response should be governed through an enterprise AI risk committee with representation from operations, AWS, legal, privacy, cybersecurity, finance, human resources, and internal audit.')
add_para(doc, 'For competitor AI advantage, Amazon should continue market scanning, benchmark AI service quality, and maintain a funded innovation backlog. This is a mitigate and compete response. For AI warehouse strategy failure, Amazon should use staged rollout gates, pilot facilities, productivity baselines, human safety reviews, and value realization tracking. For cybersecurity and data breach, the response should include zero trust access, encryption, API security testing, vendor assessment, incident response testing, and data minimization. Since this risk is a major tail contributor, it should receive executive level reporting.')
add_para(doc, 'For AI forecasting failure, Amazon should monitor forecast error, data drift, bias in demand patterns, and exception overrides. Human review should remain in place for high impact forecasts. For system integration failure, the strategy should use test environments, interface contracts, data reconciliation, rollback plans, and integration defect thresholds before scaling. For cloud service outage, the response should include multi region resilience where appropriate, backup queues, failover testing, service level monitoring, and customer impact playbooks.')
add_para(doc, 'For workforce adoption failure, Amazon should treat employees as part of the control environment rather than only as users of a new system. Training completion is not enough. Management should measure actual adoption, user confidence, escalation quality, and employee sentiment. For cost overrun and delayed ROI, finance should use stage funding, benefit ownership, variance review, and independent ROI validation. For regulatory and privacy non compliance, privacy impact assessments, data mapping, consent review, retention rules, and audit readiness should be built into the project. For AI governance and ethics failure, every high impact AI use case should have model documentation, bias testing, human oversight, and a clear owner.')
add_para(doc, 'The risk appetite should be differentiated. Amazon can accept moderate experimentation risk in low impact pilots, but should have low appetite for customer data exposure, regulatory breach, unsafe warehouse automation, and opaque AI decisions affecting customers or employees. This balanced appetite reflects COSO guidance that ERM should be connected to strategy and performance rather than treated as a separate checklist (Committee of Sponsoring Organizations of the Treadway Commission, 2017).')

add_heading(doc, 'Key Risk Indicators and Trigger Points for Action', 1)
add_para(doc, 'The KRIs below are designed as early warning signals. They are not only reporting metrics. Each trigger has a linked management action so that risk monitoring leads to a decision. The thresholds are practical starting points for a large technology and operations environment and should be refined after Amazon gathers internal baseline data.')
kri_rows = [
['Competitor AI Advantage', 'Competitor AI release index, customer service response benchmark, AI cost per contact, market share change', 'Two major competitor AI releases in one quarter or customer service benchmark gap above 10 percent', 'Trigger executive competitive review and reprioritize AI feature backlog'],
['AI Warehouse Strategy Failure', 'Warehouse productivity improvement, robot travel efficiency, robot failure rate, safety incidents', 'Productivity improvement below 5 percent after two quarters or robot failure above 3 per 1,000 tasks', 'Pause scale up, review process design, and revise pilot controls'],
['Cybersecurity & Data Breach', 'Confirmed incidents, high severity vulnerabilities, failed login attempts, mean time to detect, API findings', 'Any confirmed breach or more than 5 high severity vulnerabilities open over 14 days', 'Activate incident response, executive reporting, and remediation funding'],
['AI Forecasting Failure', 'Forecast error, model drift, stockout rate, overstock rate, override rate', 'Forecast error above 15 percent for two cycles or stockout rate above 4 percent', 'Retrain model, review data quality, and add human review for high impact forecasts'],
['System Integration Failure', 'Integration defect count, API failure rate, data sync lag, reconciliation errors', 'API failure above 2 percent daily or critical defects open over 7 days', 'Stop release, start rollback review, and deploy integration war room'],
['Cloud Service Outage', 'System uptime, downtime minutes, failed failover tests, queue backlog', 'Monthly uptime below 99.95 percent or any failed failover test for critical service', 'Run continuity plan, reroute workload, and review architecture'],
['Workforce Adoption Failure', 'Employee adoption rate, training completion, support ticket volume, employee sentiment', 'Adoption below 80 percent by target date or training completion below 90 percent', 'Launch targeted training, manager coaching, and process simplification'],
['Cost Overrun & Delayed ROI', 'Budget variance, burn rate, benefit realization, ROI forecast, scope change count', 'Budget variance above 10 percent or ROI delayed by more than one quarter', 'Require finance review, revise scope, and approve next funding stage only after correction'],
['Regulatory & Privacy Non Compliance', 'Open compliance gaps, privacy audit findings, data subject request backlog, privacy complaints', 'Any high severity privacy gap or legal deadline breach', 'Escalate to legal and privacy leadership, freeze risky processing, and remediate'],
['AI Governance & Ethics Failure', 'Model cards complete, bias reports, fairness variance, human override rate, audit exceptions', 'Model documentation below 100 percent for high impact use cases or fairness variance above 5 percent', 'Require governance review, pause deployment, and complete bias remediation'],
]
add_table(doc, ['Risk', 'KRI set', 'Trigger point', 'Required action'], kri_rows, widths=[1.25,2.3,2.1,2.0], font_size=7.2)
add_para(doc, 'The most important point is that triggers should be acted on quickly. A KRI dashboard that does not change management behavior is only a reporting tool. For this project, KRIs should be reviewed monthly by the AI risk committee, with urgent escalation for breach, outage, safety, privacy, and governance triggers. The KRI owners should also document whether each trigger was accepted, mitigated, escalated, or closed.')

add_heading(doc, 'Conclusion', 1)
add_para(doc, 'Amazon’s AI warehouse and customer service strategy has strong upside, but the enterprise risk profile is material. The qualitative assessment shows that competitor AI advantage and AI warehouse strategy failure are critical because they directly affect strategy and competitive position. The quantitative assessment adds a more rigorous view by estimating expected loss, variance, modelled tail exposure, and KRI based prediction. The base expected exposure is about 696 million dollars, while the Monte Carlo 95th percentile is about 1.59 billion dollars. This gap between average exposure and tail exposure is the strongest reason for active ERM oversight.')
add_para(doc, 'The report also shows that the largest tail risks are not only technical. Cybersecurity, regulatory non compliance, and cost overrun dominate the worst loss years. Workforce adoption and AI governance are also essential because weak adoption, poor documentation, or bias can turn a technically strong platform into a business problem. Therefore, Amazon should manage the AI program through integrated controls: security by design, privacy by design, resilient cloud architecture, staged warehouse rollout, transparent AI governance, strong training, and financial benefit tracking.')
add_para(doc, 'My main recommendation is that Amazon should proceed with AI expansion, but only with a risk response plan that is measurable and tied to decision triggers. The ten risks should remain on the enterprise risk register until implementation stabilizes. Each risk should have an executive owner, KRIs, monthly trigger review, and documented treatment actions. This approach gives Amazon a better chance to capture AI value without exposing the business to uncontrolled security, compliance, financial, and operational losses.')
add_para(doc, 'For the project repository, the recommended structure is a main enterprise risk project folder with a README file, requirements file, raw data folder, processed data folder, notebooks for exploratory analysis, qualitative analysis, quantitative analysis, and Monte Carlo analysis, plus a report folder containing the final submission. This structure supports transparency because the report, assumptions, synthetic dataset, and model outputs can be reviewed together.')

# References
doc.add_page_break()
add_heading(doc, 'References', 1)
refs = [
'Amazon. (2025). Amazon launches a new AI foundation model to power its robotic fleet and deploys its one millionth robot. About Amazon.',
'Amazon. (2026a). Amazon announces fourth quarter results. Amazon Investor Relations.',
'Amazon. (2026b). Amazon announces new robotics, faster delivery expansion in Europe, and 1 billion dollars in employee upskilling. About Amazon.',
'Amazon.com, Inc. (2026c). Annual report on Form 10 K for the fiscal year ended December 31, 2025. U.S. Securities and Exchange Commission.',
'Amazon Web Services. (2024). Reliability pillar: AWS Well Architected Framework. Amazon Web Services.',
'Committee of Sponsoring Organizations of the Treadway Commission. (2017). Enterprise risk management: Integrating with strategy and performance. COSO.',
'European Commission. (2024). AI Act: Regulatory framework for artificial intelligence. European Commission.',
'European Parliament and Council of the European Union. (2016). Regulation (EU) 2016 679 General Data Protection Regulation. Official Journal of the European Union.',
'Fleisher, C. S., & Bensoussan, B. E. (2015). Business and competitive analysis: Effective application of new and classic methods. FT Press.',
'Group 3. (2026). Assignment 1: Amazon and SWOT analysis. Unpublished course paper.',
'IBM. (2025). Cost of a data breach report 2025. IBM.',
'International Organization for Standardization. (2018). ISO 31000: Risk management guidelines. ISO.',
'National Institute of Standards and Technology. (2023). Artificial intelligence risk management framework (AI RMF 1.0). U.S. Department of Commerce.'
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(clean(ref))
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# Appendix
doc.add_page_break()
add_heading(doc, 'Appendix 1: Complete Risk Treatment and Response Plan', 1)
add_para(doc, 'The following plan translates the risk assessment into treatment actions. Owners are written by function because the final organizational owner would be assigned internally by Amazon.')
app_rows = [
['Competitor AI Advantage', 'Strategic planning and product leadership', 'Mitigate and compete', 'Maintain AI competitor intelligence, benchmark customer service quality, protect innovation budget, and review feature backlog monthly.', 'Competitor release index, response benchmark, market share change', 'If trigger occurs, hold executive review within 10 business days and approve response roadmap.'],
['AI Warehouse Strategy Failure', 'Operations and robotics leadership', 'Mitigate', 'Use pilot gates, productivity baselines, robot safety reviews, employee feedback, and independent value tracking before network rollout.', 'Productivity improvement, robot failure rate, safety incidents', 'If productivity improvement is below threshold, pause scale up and redesign process controls.'],
['Cybersecurity & Data Breach', 'Chief information security office', 'Mitigate and transfer', 'Apply zero trust access, encryption, API testing, vendor security review, data minimization, incident response exercises, and cyber insurance review.', 'Security incidents, vulnerabilities, failed logins, mean time to detect', 'If breach or high severity trigger occurs, activate incident command and notify legal and privacy leadership.'],
['AI Forecasting Failure', 'Data science and supply chain planning', 'Mitigate', 'Monitor forecast error, run drift checks, validate data quality, keep human review for high value decisions, and retrain models on schedule.', 'Forecast error, drift, stockout rate, overstock rate', 'If forecast error trigger occurs, move affected forecasts to enhanced review until model performance improves.'],
['System Integration Failure', 'Technology integration leadership', 'Mitigate', 'Create interface contracts, test APIs, reconcile data, monitor sync lag, use release gates, and maintain rollback plans.', 'Integration defects, API failures, data sync lag', 'If critical defects persist, stop release and launch integration recovery team.'],
['Cloud Service Outage', 'Cloud platform and business continuity teams', 'Mitigate and accept residual risk', 'Design resilient architecture, conduct failover tests, monitor uptime, create backup queues, and rehearse customer impact playbooks.', 'Uptime, downtime minutes, failover tests, queue backlog', 'If uptime or failover trigger occurs, execute continuity plan and complete root cause review.'],
['Workforce Adoption Failure', 'Human resources and operations training', 'Mitigate', 'Provide role based training, manager coaching, adoption support, employee communication, feedback loops, and simplified user procedures.', 'Adoption rate, training completion, sentiment, support tickets', 'If adoption is below threshold, deploy targeted retraining and remove process barriers.'],
['Cost Overrun & Delayed ROI', 'Finance and program management office', 'Mitigate', 'Use stage funding, budget variance review, benefit realization tracking, scope control, vendor cost review, and independent ROI validation.', 'Budget variance, burn rate, ROI forecast, scope changes', 'If budget or ROI trigger occurs, freeze discretionary scope and require finance approval for next stage.'],
['Regulatory & Privacy Non Compliance', 'Legal, privacy, and compliance', 'Avoid where unlawful and mitigate otherwise', 'Run privacy impact assessments, map data flows, check consent and retention rules, review cross border data issues, and maintain audit evidence.', 'Open compliance gaps, privacy findings, request backlog, complaints', 'If high severity gap appears, freeze risky processing and complete legal remediation plan.'],
['AI Governance & Ethics Failure', 'AI governance committee and internal audit', 'Mitigate', 'Require model cards, impact assessments, fairness testing, human oversight, approval records, exception tracking, and periodic audit.', 'Model documentation, bias reports, fairness variance, audit exceptions', 'If documentation or fairness trigger occurs, pause deployment until governance review is complete.'],
]
add_table(doc, ['Risk', 'Owner', 'Treatment', 'Main controls', 'KRIs', 'Response trigger action'], app_rows, widths=[1.1,1.25,1.0,2.5,1.3,2.0], font_size=6.6)

# Add footer page numbers maybe not necessary, but add simple footer text
for sec in doc.sections:
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Amazon ERM Analytics Assessment')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)

# Ensure no literal hyphen characters in all paragraphs and tables
texts = []
for p in doc.paragraphs:
    texts.append(p.text)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
alltext = '\n'.join(texts)
bad = [c for c in ['-','–','—'] if c in alltext]
if bad:
    print('Bad dash chars found', bad)
    # leave print but clean should have handled common text

# Save
doc.save(OUT)
print(OUT)
